from docx import Document

def generate_dummy_docx_file(firstName, lastName, SID, anamnesisGen, berteilung):

    document = Document()

    document.add_heading(firstName + ' ' + lastName + ' ' + SID, 0)

    document.add_paragraph('Pregenerated results')

    document.add_heading('Anamnesis', level=1)
    document.add_paragraph(str(anamnesisGen.strip()).replace('<pad>', ''), style='Intense Quote')

    document.add_heading('Berteilung', level=1)
    document.add_paragraph(str(berteilung.strip()).replace('<pad>', ''), style='Intense Quote')

    document.save('report.docx')


if __name__ == "__main__":
    generate_dummy_docx_file()